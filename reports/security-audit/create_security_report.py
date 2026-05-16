from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from pathlib import Path

OUT = Path('reports/security-audit/Ultimate_Computer_Software_Security_Audit_Report.docx')

META = {
    'project': 'Ultimate Computer Software',
    'date': 'May 10, 2026',
    'scope': 'Local source-code security review of frontend, backend, Prisma schema, routes, auth, uploads, SEO/public endpoints, dependencies, and deployment configuration.',
}

high_risks = [
    {
        'title': 'Default admin credentials remain in seed/docs',
        'severity': 'High',
        'evidence': 'backend/prisma/seed.ts:113-126; README.md:118-121',
        'risk': 'If the seeded admin account or password is still valid in production, an attacker can log in as an administrator.',
        'recommendation': 'Change the admin password immediately, remove default credentials from public docs, disable seed-admin creation in production, and require 2FA for admins.',
    },
    {
        'title': 'Admin-controlled ad HTML is rendered publicly',
        'severity': 'High',
        'evidence': 'frontend/src/components/AdBanner.tsx:59; backend/src/controllers/configController.ts:35-47',
        'risk': 'A compromised admin account or unsafe ad snippet can become stored XSS or a malicious third-party script vector across public pages.',
        'recommendation': 'Replace arbitrary HTML ads with structured ad settings, sanitize tightly, or render ads inside sandboxed iframes with a strict allowlist.',
    },
    {
        'title': 'Article body iframes allow high-risk third-party embed hosts',
        'severity': 'High',
        'evidence': 'backend/src/utils/sanitize.ts:177-189; frontend/src/pages/public/ArticlePage.tsx:85-96',
        'risk': 'Code execution platforms and broad iframe allowlists increase phishing, tracking, and content-execution risk inside article pages.',
        'recommendation': 'Allow only YouTube by default. If code demos are required, sandbox every non-YouTube iframe with minimal permissions and remove broad host regexes.',
    },
    {
        'title': 'Production protections depend heavily on NODE_ENV=production',
        'severity': 'High',
        'evidence': 'backend/src/app.ts:93-101; backend/src/middleware/rateLimiters.ts:15-25; backend/src/config/env.ts:36-69',
        'risk': 'If production is deployed with NODE_ENV=development, rate limits and CAPTCHA defaults are weaker, increasing brute-force and abuse risk.',
        'recommendation': 'Add startup checks for production deployments and fail closed unless NODE_ENV, HTTPS cookies, Turnstile, strong secrets, and proxy settings are correct.',
    },
]

medium_risks = [
    {
        'title': 'Refresh token rotation has a concurrency race',
        'severity': 'Medium',
        'evidence': 'backend/src/controllers/authController.ts:302-339; backend/src/services/refreshSessionService.ts:68-99',
        'risk': 'Two simultaneous refresh requests can both validate before revocation, reducing replay-detection strength.',
        'recommendation': 'Rotate tokens with an atomic conditional update inside one transaction and require exactly one affected row before creating the next session.',
    },
    {
        'title': 'OG image endpoint can leak unpublished article metadata',
        'severity': 'Medium',
        'evidence': 'backend/src/controllers/ogController.ts:145-159',
        'risk': 'A known draft/submitted slug can reveal title, author, and category through /og/article/:slug.',
        'recommendation': 'Require status=PUBLISHED for public OG image generation unless caller is authenticated as owner/admin.',
    },
    {
        'title': 'Public series endpoint can expose draft article metadata',
        'severity': 'Medium',
        'evidence': 'backend/src/controllers/articleController.ts:822-843',
        'risk': 'Series responses include member article status, title, slug, excerpt, and image even if articles are not published.',
        'recommendation': 'Filter public series members to published articles. Show unpublished members only to owner/admin.',
    },
    {
        'title': 'Poll endpoints can expose unpublished article polls',
        'severity': 'Medium',
        'evidence': 'backend/src/controllers/pollController.ts:89-124 and 220-255',
        'risk': 'Polls linked to unpublished content may be retrievable by ID or article ID.',
        'recommendation': 'When a poll has articleId, verify linked article visibility before returning poll data.',
    },
    {
        'title': 'Article view endpoint can be abused to inflate reads',
        'severity': 'Medium',
        'evidence': 'backend/src/controllers/articleController.ts:497-575',
        'risk': 'Anonymous clients can repeatedly POST view events, inflating stats and causing database write load.',
        'recommendation': 'Add IP/user rate limits, anonymous dedupe window, capped timeSpentSeconds, and optional bot/captcha protection for suspicious bursts.',
    },
    {
        'title': '2FA recovery codes are shorter than recommended',
        'severity': 'Medium',
        'evidence': 'backend/src/utils/totp.ts:101-110',
        'risk': 'Current recovery codes are 32-bit each, which is weaker than typical backup-code entropy.',
        'recommendation': 'Generate stronger one-time recovery codes, for example 10 to 16 random bytes per code, and keep only hashes in the database.',
    },
    {
        'title': 'Disabling 2FA does not revoke active refresh sessions',
        'severity': 'Medium',
        'evidence': 'backend/src/controllers/authController.ts:481-534',
        'risk': 'Existing sessions remain valid after a major security setting changes.',
        'recommendation': 'Revoke all refresh sessions after enabling/disabling 2FA, password reset, role changes, bans, and other high-risk account actions.',
    },
    {
        'title': 'Dormant data processing functions could be dangerous if exposed later',
        'severity': 'Medium',
        'evidence': 'backend/src/services/dataProcessing.ts:13-257',
        'risk': 'The service can export or deactivate accounts based only on email if a future route calls it without strong verification.',
        'recommendation': 'Remove unused dangerous functions or require authenticated/verified identity plus admin approval and audit logging.',
    },
]

low_risks = [
    {
        'title': 'Frontend dependency advisory for Quill',
        'severity': 'Low',
        'evidence': 'npm audit --omit=dev in frontend reports quill XSS advisory through react-quill-new',
        'risk': 'Rich text editor dependency has a known low-severity advisory.',
        'recommendation': 'Track upstream fix, update when compatible, and keep server-side article sanitization as the primary defense.',
    },
    {
        'title': 'Frontend public robots.txt allows all routes',
        'severity': 'Low',
        'evidence': 'frontend/public/robots.txt',
        'risk': 'If the frontend static robots file overrides the backend robots endpoint, admin/private routes are not disallowed for crawlers.',
        'recommendation': 'Align frontend robots.txt with backend robots rules and disallow /admin, /dashboard, /login, /register, /reset-password, /api/admin, and /api/me.',
    },
    {
        'title': 'Accidental empty .crdownload file in public folder',
        'severity': 'Low',
        'evidence': 'frontend/public/Unconfirmed 716710.crdownload',
        'risk': 'The file is empty, but public folders should not contain accidental browser download artifacts.',
        'recommendation': 'Delete accidental public files and add ignore rules/checks to keep them out of builds.',
    },
    {
        'title': 'Upload pre-filter trusts client MIME type',
        'severity': 'Low',
        'evidence': 'backend/src/middleware/upload.ts:24-35; backend/src/controllers/uploadController.ts:23-33',
        'risk': 'Sharp rejects invalid images later, but MIME type alone is client-controlled and should not be the first trust boundary.',
        'recommendation': 'Validate magic bytes and Sharp metadata before processing; add pixel-count limits to reduce decompression-bomb risk.',
    },
]

strong_controls = [
    'Authentication uses HTTP-only cookies plus CSRF protection.',
    'Access-token verification re-checks the user and role from the database.',
    'Refresh tokens are stored as hashes in database-backed sessions.',
    'Password reset tokens are hashed, expiring, and single-use.',
    'Prisma is used for database access; no practical raw SQL injection pattern was found in active code.',
    'Admin routes are protected by backend authentication and role middleware.',
    'Public contact/data requests are queued instead of automatically exporting or deleting accounts.',
    'Article body HTML is sanitized server-side and again client-side before rendering.',
]

priority_fixes = [
    'Change and rotate the admin password; remove default credentials from docs and seed output.',
    'Require 2FA for admin and moderator accounts before production launch.',
    'Replace raw ad HTML with safe structured configuration or sandboxed iframe rendering.',
    'Restrict article embeds to YouTube or sandbox all non-YouTube embeds with minimal permissions.',
    'Fail startup when production security environment variables are missing or unsafe.',
    'Fix public metadata leaks in OG images, series, polls, pinned profile articles, and notification hydration.',
    'Make refresh-token rotation atomic and revoke sessions after security changes.',
    'Add abuse limits to views, comments, reactions, follows, polls, and article-request endpoints.',
    'Clean public/static folders and align robots.txt with backend rules.',
]

prod_checklist = [
    'NODE_ENV=production',
    'HTTPS enforced at the proxy/CDN',
    'COOKIE_SECURE=true',
    'COOKIE_SAME_SITE=lax or strict unless cross-site auth is required',
    'JWT_SECRET and JWT_REFRESH_SECRET generated with high entropy and never committed',
    'TWO_FACTOR_ENCRYPTION_KEY set separately from JWT secrets',
    'CAPTCHA_REQUIRED=true and CAPTCHA_SECRET configured',
    'TRUST_PROXY correctly set behind the production reverse proxy',
    'Admin and moderator accounts protected by 2FA',
    'Database backups encrypted and access restricted to production operators only',
    'Centralized logs enabled for auth failures, admin actions, password resets, bans, and report status changes',
]

commands = [
    'npm audit --omit=dev',
    'rg "dangerouslySetInnerHTML|eval|localStorage|document.cookie|child_process|$queryRaw"',
    'rg "TODO|FIXME|backdoor|bypass|Admin@12345"',
    'find . -name ".env*" -print',
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in('w:tcW')
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(int(width_inches * 1440)))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_borders(table, color='DADCE0', size='6'):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = 'w:' + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def set_table_width(table, width_inches=6.5):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(int(width_inches * 1440)))
    tbl_w.set(qn('w:type'), 'dxa')


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)

def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)

def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in('w:tblCellMar')
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement('w:tblCellMar')
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tbl_cell_mar.find(qn(f'w:{margin_name}'))
        if node is None:
            node = OxmlElement(f'w:{margin_name}')
            tbl_cell_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    set_paragraph_spacing(p, before=0, after=4, line=1.15)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    set_paragraph_spacing(p, before=0, after=4, line=1.15)
    return p


def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + ': ')
    r.bold = True
    p.add_run(text)
    set_paragraph_spacing(p, before=0, after=4, line=1.15)
    return p


def add_risk_detail(doc, item):
    doc.add_heading(item['title'], level=3)
    add_label_para(doc, 'Severity', item['severity'])
    add_label_para(doc, 'Evidence', item['evidence'])
    add_label_para(doc, 'Risk', item['risk'])
    add_label_para(doc, 'Recommendation', item['recommendation'])


def add_risk_table(doc, items):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 6.5)
    set_table_borders(table)
    set_cell_margins(table)
    repeat_table_header(table.rows[0])
    keep_row_together(table.rows[0])
    headers = ['Severity', 'Finding', 'Primary action']
    widths = [0.8, 3.45, 2.25]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        for p in cell.paragraphs:
            set_paragraph_spacing(p, before=0, after=0, line=1.15)
    for item in items:
        row = table.add_row()
        keep_row_together(row)
        cells = row.cells
        if item['severity'] == 'High':
            primary = 'Pre-launch fix.'
        elif item['severity'] == 'Medium':
            primary = 'Next hardening pass.'
        else:
            primary = 'Routine cleanup.'
        values = [item['severity'], item['title'], primary]
        for idx, value in enumerate(values):
            set_cell_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[idx].paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(10)
            set_paragraph_spacing(p, before=0, after=0, line=1.15)
    return table


def style_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    normal.font.size = Pt(11)
    normal.font.color.rgb = None
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title = styles['Title']
    title.font.name = 'Arial'
    title._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    title.font.size = Pt(24)
    title.font.bold = True
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.15

    for name, size, before, after in [
        ('Heading 1', 16, 18, 8),
        ('Heading 2', 14, 14, 6),
        ('Heading 3', 12, 10, 4),
    ]:
        st = styles[name]
        st.font.name = 'Arial'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = None
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15

    for name in ['List Bullet', 'List Number']:
        if name in styles:
            st = styles[name]
            st.font.name = 'Arial'
            st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            st.font.size = Pt(11)
            st.paragraph_format.space_after = Pt(4)
            st.paragraph_format.line_spacing = 1.15


def build():
    doc = Document()
    style_document(doc)
    props = doc.core_properties
    props.title = 'Ultimate Computer Software Security Audit Report'
    props.subject = 'Security risks and recommendations'
    props.author = 'OpenAI Codex'
    props.comments = 'Generated from local source-code audit. No secrets from .env are included.'

    p = doc.add_paragraph(style='Title')
    p.add_run('Ultimate Computer Software Security Audit Report')
    subtitle = doc.add_paragraph()
    subtitle.add_run('Risks, evidence, and hardening recommendations').italic = True
    set_paragraph_spacing(subtitle, before=0, after=8, line=1.15)
    add_label_para(doc, 'Report date', META['date'])
    add_label_para(doc, 'Scope', META['scope'])
    add_label_para(doc, 'Important limitation', 'This is a source-code security review, not a full penetration test. Findings should be validated again in the deployed production environment.')

    doc.add_heading('Executive Summary', level=1)
    add_bullet(doc, 'No obvious deliberate public backdoor was found, such as an unauthenticated admin route, hardcoded login bypass, database dump endpoint, or active raw SQL injection path.')
    add_bullet(doc, 'The highest practical risks are default admin credentials, arbitrary admin-managed ad HTML, broad iframe/embed permissions, and production environment misconfiguration.')
    add_bullet(doc, 'The application already has several strong controls: HTTP-only auth cookies, CSRF protection, hashed refresh sessions, hashed reset tokens, Prisma query safety, and backend admin role checks.')
    add_bullet(doc, 'The recommended launch posture is to fix the high-risk items first, then address metadata leaks, session hardening, and abuse-rate limits before public production use.')

    doc.add_page_break()
    doc.add_heading('Risk Register', level=1)
    doc.add_paragraph('The table below summarizes all identified risks and the primary recommended action for each.')
    add_risk_table(doc, high_risks + medium_risks + low_risks)

    doc.add_heading('Detailed Findings', level=1)
    doc.add_heading('High Risk', level=2)
    for item in high_risks:
        add_risk_detail(doc, item)

    doc.add_heading('Medium Risk', level=2)
    for item in medium_risks:
        add_risk_detail(doc, item)

    doc.add_heading('Low Risk', level=2)
    for item in low_risks:
        add_risk_detail(doc, item)

    doc.add_heading('Strong Controls Already Present', level=1)
    for item in strong_controls:
        add_bullet(doc, item)

    doc.add_heading('Priority Remediation Plan', level=1)
    for item in priority_fixes:
        add_number(doc, item)

    doc.add_heading('Production Launch Security Checklist', level=1)
    for item in prod_checklist:
        add_bullet(doc, item)

    doc.add_heading('Audit Commands Used', level=1)
    p = doc.add_paragraph('Representative local checks used during the review:')
    set_paragraph_spacing(p)
    for cmd in commands:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(cmd)
        run.font.name = 'Courier New'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        set_paragraph_spacing(p, before=0, after=4, line=1.15)


    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    build()
